"""
Document Text Extraction Module

This module provides text extraction for Microsoft Word documents (.docx format)
using the python-docx library.

Classes:
    DocxExtractor: DOCX text extractor with sampling for large documents
"""

import docx
from pathlib import Path
from typing import Union, Optional

from .base_extractor import BaseExtractor


class DocxExtractor(BaseExtractor):
    """
    Text extractor for DOCX documents.

    Extracts text from Word documents using python-docx. For large documents
    (>180 paragraphs), samples first and last paragraphs to limit processing time.
    Also extracts text from tables.

    Attributes:
        PARAGRAPH_LIMIT_FOR_SAMPLING (int): Paragraph threshold for sampling (180)
        PARAGRAPHS_TO_SAMPLE (int): Paragraphs to extract from start/end when sampling (90)
    """
    def __init__(self):
        super().__init__()

        # Sampling settings for large files
        self.PARAGRAPH_LIMIT_FOR_SAMPLING = 180
        self.PARAGRAPHS_TO_SAMPLE = 90

    def extract(self, input_path: Union[str, Path]) -> Optional[str]:
        """
        Extract text from DOCX file including tables.

        Args:
            input_path: Path to DOCX file

        Returns:
            Optional[str]: Extracted text content if successful, None if extraction failed.
        """
        docx_path = Path(input_path)
        source_filename = docx_path.name

        # Validate file using inherited method
        validation_error = self._validate_file(docx_path, '.docx')
        if validation_error:
            self.logger.error(f"DOCX validation failed for '{source_filename}': {validation_error}")
            return None

        try:
            # Extract text using internal method
            extracted_text = self._extract_text(docx_path, source_filename)
            return extracted_text

        except Exception as e:
            self.logger.error(f"DOCX extraction failed for '{source_filename}': {e}")
            return None

    def _extract_text(self, docx_path: Path, source_filename: str) -> str:
        """
        Extract text with sampling for large documents.

        Strategy:
            - ≤180 paragraphs: Extract all paragraphs + tables
            - >180 paragraphs: Extract first 90 + last 90 paragraphs only
            - Empty paragraphs are filtered out

        Args:
            docx_path: Path to DOCX file
            source_filename: Filename for logging

        Returns:
            str: Extracted text with paragraphs joined by double newlines
        """
        document = docx.Document(str(docx_path))
        paragraphs = document.paragraphs
        num_paragraphs = len(paragraphs)

        full_text_parts = []

        # Sampling logic for large documents
        if num_paragraphs > self.PARAGRAPH_LIMIT_FOR_SAMPLING:
            self.logger.info(
                f"'{source_filename}' has {num_paragraphs} paragraphs. "
                f"Sampling the first {self.PARAGRAPHS_TO_SAMPLE} and last {self.PARAGRAPHS_TO_SAMPLE}."
            )

            # Extract first paragraphs
            for i in range(min(self.PARAGRAPHS_TO_SAMPLE, len(paragraphs))):
                if paragraphs[i].text.strip():
                    full_text_parts.append(paragraphs[i].text)

            # Add separator
            full_text_parts.append("\n\n... (intermediate paragraph content omitted) ...\n\n")

            # Extract last paragraphs
            start_last_paragraphs = max(0, num_paragraphs - self.PARAGRAPHS_TO_SAMPLE)
            for i in range(start_last_paragraphs, num_paragraphs):
                if paragraphs[i].text.strip():
                    full_text_parts.append(paragraphs[i].text)

        else:
            self.logger.info(f"'{source_filename}' has {num_paragraphs} paragraphs. Extracting all content.")

            # Extract text from all paragraphs
            for para in paragraphs:
                if para.text.strip():
                    full_text_parts.append(para.text)

            # Extract text from all tables
            if document.tables:
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text_parts.append(cell.text)

        extracted_text = "\n\n".join(full_text_parts)
        self.logger.info(f"Extracted {len(extracted_text)} characters from '{source_filename}'")

        return extracted_text

